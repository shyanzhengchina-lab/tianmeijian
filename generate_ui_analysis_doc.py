#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成MES系统UI风格分析WORD文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_ui_analysis_doc():
    """创建UI风格分析文档"""

    # 创建文档
    doc = Document()

    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    # 添加标题
    title = doc.add_heading('MES系统前端UI风格和视觉风格分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加分隔线
    doc.add_paragraph('_' * 80)

    # 1. 设计系统概述
    doc.add_heading('1. 设计系统', level=1)
    p1 = doc.add_paragraph()
    p1.add_run('UI框架: ').bold = True
    p1.add_run('Ant Design (antd) - 企业级React组件库')
    p2 = doc.add_paragraph()
    p2.add_run('设计语言: ').bold = True
    p2.add_run('企业级B端管理系统风格')
    p3 = doc.add_paragraph()
    p3.add_run('布局方式: ').bold = True
    p3.add_run('Flexbox + CSS Grid')

    # 2. 颜色系统
    doc.add_heading('2. 颜色系统', level=1)

    # 2.1 主色调
    doc.add_heading('2.1 主色调', level=2)
    table_main_colors = doc.add_table(rows=5, cols=3)
    table_main_colors.style = 'Table Grid'
    table_main_colors.rows[0].cells[0].text = '颜色名称'
    table_main_colors.rows[0].cells[1].text = '颜色值'
    table_main_colors.rows[0].cells[2].text = '用途说明'

    table_main_colors.rows[1].cells[0].text = '品牌红'
    table_main_colors.rows[1].cells[1].text = '#C8000A'
    table_main_colors.rows[1].cells[2].text = '顶部导航栏（突出企业品牌）'

    table_main_colors.rows[2].cells[0].text = '主题蓝'
    table_main_colors.rows[2].cells[1].text = '#1677ff'
    table_main_colors.rows[2].cells[2].text = '主交互色、选中状态、链接'

    table_main_colors.rows[3].cells[0].text = '背景灰'
    table_main_colors.rows[3].cells[1].text = '#F0F2F5'
    table_main_colors.rows[3].cells[2].text = '主背景色'

    table_main_colors.rows[4].cells[0].text = '白色'
    table_main_colors.rows[4].cells[1].text = '#ffffff'
    table_main_colors.rows[4].cells[2].text = '卡片/模块背景'

    # 2.2 语义色彩
    doc.add_heading('2.2 语义色彩（状态编码）', level=2)
    table_semantic = doc.add_table(rows=6, cols=4)
    table_semantic.style = 'Table Grid'
    table_semantic.rows[0].cells[0].text = '状态'
    table_semantic.rows[0].cells[1].text = '颜色'
    table_semantic.rows[0].cells[2].text = '用途'
    table_semantic.rows[0].cells[3].text = '十六进制'

    table_semantic.rows[1].cells[0].text = '✅ 成功'
    table_semantic.rows[1].cells[1].text = '绿色'
    table_semantic.rows[1].cells[2].text = '生产中、完成'
    table_semantic.rows[1].cells[3].text = '#52c41a'

    table_semantic.rows[2].cells[0].text = '⚠️ 警告'
    table_semantic.rows[2].cells[1].text = '黄色'
    table_semantic.rows[2].cells[2].text = '待转移、预警'
    table_semantic.rows[2].cells[3].text = '#faad14'

    table_semantic.rows[3].cells[0].text = 'ℹ️ 信息'
    table_semantic.rows[3].cells[1].text = '蓝色'
    table_semantic.rows[3].cells[2].text = '待检验、处理中'
    table_semantic.rows[3].cells[3].text = '#1677ff'

    table_semantic.rows[4].cells[0].text = '❌ 错误'
    table_semantic.rows[4].cells[1].text = '红色'
    table_semantic.rows[4].cells[2].text = '异常停机、阻塞'
    table_semantic.rows[4].cells[3].text = '#ff4d4f'

    table_semantic.rows[5].cells[0].text = '⚪ 默认'
    table_semantic.rows[5].cells[1].text = '灰色'
    table_semantic.rows[5].cells[2].text = '未开工、闲置'
    table_semantic.rows[5].cells[3].text = '#8c8c8c'

    # 2.3 文字层次
    doc.add_heading('2.3 文字层次', level=2)
    table_text = doc.add_table(rows=4, cols=2)
    table_text.style = 'Table Grid'
    table_text.rows[0].cells[0].text = '文字类型'
    table_text.rows[0].cells[1].text = '颜色值'

    table_text.rows[1].cells[0].text = '主文字'
    table_text.rows[1].cells[1].text = '#1A1A1A / #262626'

    table_text.rows[2].cells[0].text = '次要文字'
    table_text.rows[2].cells[1].text = '#8C8C8C'

    table_text.rows[3].cells[0].text = '辅助文字'
    table_text.rows[3].cells[1].text = '#98A2B3 / #667085'

    # 3. 字体系统
    doc.add_heading('3. 字体系统', level=1)
    p4 = doc.add_paragraph()
    p4.add_run('字体栈: ').bold = True
    p4.add_run('-apple-system, BlinkMacSystemFont, \'PingFang SC\', \'Microsoft YaHei\', \'Segoe UI\'')

    doc.add_heading('字号层级', level=2)
    table_font = doc.add_table(rows=4, cols=3)
    table_font.style = 'Table Grid'
    table_font.rows[0].cells[0].text = '类型'
    table_font.rows[0].cells[1].text = '字号'
    table_font.rows[0].cells[2].text = '字重'

    table_font.rows[1].cells[0].text = '标题'
    table_font.rows[1].cells[1].text = '16-26px'
    table_font.rows[1].cells[2].text = 'Font-weight: 600-800'

    table_font.rows[2].cells[0].text = '正文'
    table_font.rows[2].cells[1].text = '12-14px'
    table_font.rows[2].cells[2].text = 'Font-weight: 400-500'

    table_font.rows[3].cells[0].text = '辅助'
    table_font.rows[3].cells[1].text = '10-11px'
    table_font.rows[3].cells[2].text = 'Font-weight: 400-600'

    p5 = doc.add_paragraph()
    p5.add_run('等宽字体: ').bold = True
    p5.add_run('Menlo, Consolas - 用于工单号、时间戳等')

    # 4. 布局结构
    doc.add_heading('4. 布局结构', level=1)
    layout_code = '''
┌─────────────────────────────────────┐
│  顶部导航 (50px, 红色)              │
├──────────┬──────────────────────────┤
│          │ 面包屑导航               │
│  侧边栏  ├──────────────────────────┤
│ (白色)   │ 主内容区 (灰色背景)      │
│          │  - KPI卡片               │
│          │  - 数据表格              │
│          │  - 操作栏                │
├──────────┴──────────────────────────┤
│  底部Tab栏 (50px, 白色)            │
└─────────────────────────────────────┘
'''
    p6 = doc.add_paragraph()
    p6.add_run(layout_code).font.name = 'Courier New'

    # 5. 组件风格
    doc.add_heading('5. 组件风格', level=1)

    # 5.1 卡片式设计
    doc.add_heading('5.1 卡片式设计', level=2)
    table_card = doc.add_table(rows=5, cols=2)
    table_card.style = 'Table Grid'
    table_card.rows[0].cells[0].text = '属性'
    table_card.rows[0].cells[1].text = '值'

    table_card.rows[1].cells[0].text = '圆角'
    table_card.rows[1].cells[1].text = '8-10px'

    table_card.rows[2].cells[0].text = '阴影'
    table_card.rows[2].cells[1].text = '0 1px 6px rgba(0,0,0,0.07) / 悬停时 0 6px 18px'

    table_card.rows[3].cells[0].text = '边框'
    table_card.rows[3].cells[1].text = '1px solid #f0f0f0'

    table_card.rows[4].cells[0].text = '背景'
    table_card.rows[4].cells[1].text = '白色 #ffffff'

    # 5.2 按钮样式
    doc.add_heading('5.2 按钮样式', level=2)
    table_button = doc.add_table(rows=4, cols=2)
    table_button.style = 'Table Grid'
    table_button.rows[0].cells[0].text = '属性'
    table_button.rows[0].cells[1].text = '值'

    table_button.rows[1].cells[0].text = '尺寸'
    table_button.rows[1].cells[1].text = '小号(28px) / 中号(32px) / 大号(40px)'

    table_button.rows[2].cells[0].text = '圆角'
    table_button.rows[2].cells[1].text = '4-6px'

    table_button.rows[3].cells[0].text = '交互'
    table_button.rows[3].cells[1].text = '悬停时边框变蓝、文字变蓝'

    # 5.3 表格设计
    doc.add_heading('5.3 表格设计', level=2)
    table_table = doc.add_table(rows=5, cols=2)
    table_table.style = 'Table Grid'
    table_table.rows[0].cells[0].text = '元素'
    table_table.rows[0].cells[1].text = '样式'

    table_table.rows[1].cells[0].text = '头部'
    table_table.rows[1].cells[1].text = '浅灰背景 #fafafa，字体加粗'

    table_table.rows[2].cells[0].text = '行悬停'
    table_table.rows[2].cells[1].text = '#fafafa'

    table_table.rows[3].cells[0].text = '选中行'
    table_table.rows[3].cells[1].text = '#e6f7ff'

    table_table.rows[4].cells[0].text = '分页'
    table_table.rows[4].cells[1].text = '右对齐，16px间距'

    # 6. 交互特点
    doc.add_heading('6. 交互特点', level=1)

    # 6.1 动画效果
    doc.add_heading('6.1 动画效果', level=2)
    table_anim = doc.add_table(rows=3, cols=2)
    table_anim.style = 'Table Grid'
    table_anim.rows[0].cells[0].text = '属性'
    table_anim.rows[0].cells[1].text = '值'

    table_anim.rows[1].cells[0].text = '过渡时间'
    table_anim.rows[1].cells[1].text = '0.15s - 0.3s'

    table_anim.rows[2].cells[0].text = '缓动函数'
    table_anim.rows[2].cells[1].text = 'ease / cubic-bezier'

    p7 = doc.add_paragraph()
    p7.add_run('悬停效果: ').bold = True
    p7.add_run('卡片上浮2px、阴影加深')

    # 6.2 状态反馈
    doc.add_heading('6.2 状态反馈', level=2)
    table_feedback = doc.add_table(rows=3, cols=2)
    table_feedback.style = 'Table Grid'
    table_feedback.rows[0].cells[0].text = '状态类型'
    table_feedback.rows[0].cells[1].text = '视觉效果'

    table_feedback.rows[1].cells[0].text = '异常状态'
    table_feedback.rows[1].cells[1].text = '红色边框闪烁动画 (1.2s infinite)'

    table_feedback.rows[2].cells[0].text = '进行中'
    table_feedback.rows[2].cells[1].text = '蓝色发光脉冲效果'

    p8 = doc.add_paragraph()
    p8.add_run('选中项: ').bold = True
    p8.add_run('左侧蓝色边框 + 蓝色勾号 ✓')

    # 6.3 滚动条
    doc.add_heading('6.3 滚动条', level=2)
    table_scroll = doc.add_table(rows=3, cols=2)
    table_scroll.style = 'Table Grid'
    table_scroll.rows[0].cells[0].text = '属性'
    table_scroll.rows[0].cells[1].text = '值'

    table_scroll.rows[1].cells[0].text = '宽度'
    table_scroll.rows[1].cells[1].text = '4-6px'

    table_scroll.rows[2].cells[0].text = '颜色'
    table_scroll.rows[2].cells[1].text = '滑块 #D9D9D9，悬停时 #BFBFBF'

    # 7. 视觉特征
    doc.add_heading('7. 视觉特征', level=1)

    # 7.1 设计理念
    doc.add_heading('7.1 设计理念', level=2)
    design_points = [
        ('卡片化', '信息模块清晰分隔'),
        ('状态驱动', '颜色编码直观传达设备/工单状态'),
        ('信息密度', '适中，关键信息突出'),
        ('专业感', '规整的网格布局、精确的间距')
    ]

    for point, desc in design_points:
        p = doc.add_paragraph()
        p.add_run(f'• {point}: ').bold = True
        p.add_run(desc)

    # 7.2 特色元素
    doc.add_heading('7.2 特色元素', level=2)
    features = [
        ('实时看板', '带脉冲动画的状态指示灯'),
        ('进度条', '渐变色填充，带百分比标签'),
        ('徽章', '圆角pill标签，带语义色背景'),
        ('抽屉', '侧滑详情面板，保持上下文')
    ]

    for feature, desc in features:
        p = doc.add_paragraph()
        p.add_run(f'• {feature}: ').bold = True
        p.add_run(desc)

    # 8. 主题支持
    doc.add_heading('8. 主题支持', level=1)

    # 8.1 亮色模式
    doc.add_heading('8.1 亮色模式（默认）', level=2)
    table_light = doc.add_table(rows=4, cols=2)
    table_light.style = 'Table Grid'
    table_light.rows[0].cells[0].text = '元素类型'
    table_light.rows[0].cells[1].text = '颜色值'

    table_light.rows[1].cells[0].text = '背景'
    table_light.rows[1].cells[1].text = '#F0F2F5 / #ffffff'

    table_light.rows[2].cells[0].text = '文字'
    table_light.rows[2].cells[1].text = '#1A1A1A / #262626'

    table_light.rows[3].cells[0].text = '边框'
    table_light.rows[3].cells[1].text = '#F0F0F0 / #E8E8E8'

    # 8.2 暗色模式
    doc.add_heading('8.2 暗色模式', level=2)
    table_dark = doc.add_table(rows=4, cols=2)
    table_dark.style = 'Table Grid'
    table_dark.rows[0].cells[0].text = '元素类型'
    table_dark.rows[0].cells[1].text = '颜色值'

    table_dark.rows[1].cells[0].text = '背景'
    table_dark.rows[1].cells[1].text = '#1f1f1f / #262626'

    table_dark.rows[2].cells[0].text = '文字'
    table_dark.rows[2].cells[1].text = '#d9d9d9 / #bfbfbf'

    table_dark.rows[3].cells[0].text = '边框'
    table_dark.rows[3].cells[1].text = '#303030 / #434343'

    # 9. 适用场景
    doc.add_heading('9. 适用场景', level=1)
    scenarios = [
        ('✅ 生产管理系统', '状态实时监控'),
        ('✅ 数据密集型应用', '列表、表格、看板'),
        ('✅ 企业后台系统', '规范、专业、易用'),
        ('✅ 移动端适配', '底部Tab栏、响应式布局')
    ]

    for scenario, desc in scenarios:
        p = doc.add_paragraph()
        p.add_run(f'{scenario}: ').bold = True
        p.add_run(desc)

    # 10. 总结
    doc.add_heading('10. 设计总结', level=1)
    summary = '''
该MES系统UI设计采用现代化企业级设计语言，具有以下特点：

1. 清晰的信息层次：通过卡片化布局和合理的间距，使信息结构清晰易读。

2. 强状态语义：使用5种标准色彩（绿/黄/蓝/红/灰）系统化表达设备和工单状态，
   便于生产人员快速识别。

3. 专业工业风格：配色稳重，符合制造业应用的严谨性，同时保持视觉舒适度。

4. 高效交互设计：悬停反馈、状态动画、选中标记等微交互增强用户体验。

5. 良好的响应式支持：适配桌面端和移动端，底部Tab栏设计适合触屏操作。

整体评价：这是一个清晰、高效、专业的企业级MES系统界面设计，状态语义化强，
适合工业生产场景的实时监控和数据展示需求。
'''
    p9 = doc.add_paragraph(summary)

    # 添加页脚
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.text = '文档生成时间: 2026年5月3日'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 保存文档
    output_path = 'C:/NEWMES/deca/MES系统UI风格分析报告.docx'
    doc.save(output_path)
    print(f"文档已生成: {output_path}")
    return output_path

if __name__ == '__main__':
    create_ui_analysis_doc()
